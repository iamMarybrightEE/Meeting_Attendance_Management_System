import os
import shutil
import stat
import sys
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.entities import (
    ChatMessage,
    ChatSession,
    IndexStatus,
    Meeting,
    MessageRole,
    Transcript,
    TranscriptChunk,
)


def ensure_storage_dir() -> Path:
    path = Path(settings.storage_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_minutes_file(meeting_id: str, filename: str, data: bytes) -> str:
    storage = ensure_storage_dir()
    safe_name = filename.replace("/", "_")
    target = storage / f"{meeting_id}_minutes_{uuid.uuid4().hex[:8]}_{safe_name}"
    target.write_bytes(data)
    return str(target)


def extract_minutes_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    raise ValueError("Unsupported format; upload PDF or DOCX only.")


def chroma_dir_for_meeting(meeting_id: str) -> Path:
    if settings.chroma_persist_root:
        return Path(settings.chroma_persist_root) / meeting_id
    return ensure_storage_dir() / "chroma" / meeting_id


def latest_chroma_dir_for_meeting(meeting_id: str) -> Path:
    """
    Resolve the directory used for retrieval.
    Supports both legacy single-dir layout and new per-run subdirectories.
    """
    base = chroma_dir_for_meeting(meeting_id)
    if not base.exists():
        return base

    direct_db = base / "chroma.sqlite3"
    if direct_db.exists():
        return base

    run_dirs = [p for p in base.iterdir() if p.is_dir()]
    if not run_dirs:
        return base
    run_dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return run_dirs[0]


def _chmod_u_rw(path: Path) -> None:
    try:
        mode = path.stat().st_mode
        os.chmod(path, mode | stat.S_IWRITE | stat.S_IREAD)
    except OSError:
        pass


def _rmtree_onexc(func, path: str, exc: BaseException) -> None:
    try:
        _chmod_u_rw(Path(path))
        func(path)
    except OSError:
        raise exc


def _rmtree_onerror(func, path: str, exc_info) -> None:
    try:
        _chmod_u_rw(Path(path))
        func(path)
    except OSError:
        pass


def clear_meeting_chroma(meeting_id: str) -> None:
    chroma_path = chroma_dir_for_meeting(meeting_id)
    if chroma_path.exists():
        if sys.version_info >= (3, 12):
            shutil.rmtree(chroma_path, onexc=_rmtree_onexc)
        else:
            shutil.rmtree(chroma_path, onerror=_rmtree_onerror)
    chroma_path.mkdir(parents=True, exist_ok=True)


def _assert_dir_writable(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)
    probe = path / ".write_probe_delete_me"
    try:
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except OSError as exc:
        raise RuntimeError(
            f"Directory is not writable: {path}. "
            "Fix ownership/permissions on meeting_agent_service/storage (or set MEETING_AGENT_STORAGE_DIR "
            "to a folder your user owns, e.g. /tmp/mams_agent_data)."
        ) from exc


def create_or_update_meeting(
    db: Session,
    *,
    meeting_id: str,
    tenant_id: str,
    title: str,
    organizer_id: str | None,
    recording_uri: str | None,
) -> Meeting:
    meeting = db.get(Meeting, meeting_id)
    if not meeting:
        meeting = Meeting(id=meeting_id, tenant_id=tenant_id, title=title)
        db.add(meeting)
    meeting.organizer_id = organizer_id
    meeting.recording_uri = recording_uri
    meeting.index_status = IndexStatus.pending
    db.commit()
    db.refresh(meeting)
    return meeting


def split_into_chunks(text: str, chunk_size: int, overlap: int) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_text(text)


def run_indexing_job(db: Session, meeting_id: str) -> int:
    meeting = db.get(Meeting, meeting_id)
    if not meeting:
        raise ValueError("meeting not found")
    transcript = db.execute(
        select(Transcript).where(Transcript.meeting_id == meeting_id).order_by(Transcript.id.desc())
    ).scalars().first()
    if not transcript:
        raise ValueError("transcript not found")
    if not transcript.full_text.strip():
        raise ValueError("transcript is empty")

    meeting.index_status = IndexStatus.processing
    db.commit()

    db.execute(delete(TranscriptChunk).where(TranscriptChunk.meeting_id == meeting_id))
    db.commit()

    base_chroma_path = chroma_dir_for_meeting(meeting_id).resolve()
    base_chroma_path.mkdir(parents=True, exist_ok=True)
    _assert_dir_writable(base_chroma_path)

    chunk_texts = split_into_chunks(transcript.full_text, settings.chunk_size, settings.chunk_overlap)
    for idx, chunk in enumerate(chunk_texts):
        db.add(
            TranscriptChunk(
                meeting_id=meeting_id,
                transcript_id=transcript.id,
                chunk_index=idx,
                text=chunk,
                vector_ref=f"chroma-{meeting_id}-{idx}",
            )
        )
    db.commit()

    rows = db.execute(
        select(TranscriptChunk)
        .where(TranscriptChunk.meeting_id == meeting_id)
        .order_by(TranscriptChunk.chunk_index.asc())
    ).scalars().all()

    documents = [
        Document(
            page_content=c.text,
            metadata={"chunk_id": c.id, "chunk_index": c.chunk_index},
        )
        for c in rows
    ]

    run_chroma_path = base_chroma_path / f"run_{uuid.uuid4().hex[:8]}"
    run_chroma_path.mkdir(parents=True, exist_ok=True)

    embedding = OllamaEmbeddings(
        model=settings.ollama_embed_model,
        base_url=settings.ollama_base_url,
    )

    try:
        Chroma.from_documents(
            documents=documents,
            embedding=embedding,
            persist_directory=str(run_chroma_path),
        )
    except Exception:
        meeting.index_status = IndexStatus.failed
        db.commit()
        raise

    meeting.index_status = IndexStatus.ready
    db.commit()
    return len(rows)


def create_chat_session(db: Session, meeting_id: str, tenant_id: str, user_id: str) -> ChatSession:
    session = ChatSession(id=uuid.uuid4().hex, meeting_id=meeting_id, tenant_id=tenant_id, user_id=user_id)
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


MEETING_PROMPT = ChatPromptTemplate.from_template(
    """You are an assistant helping users understand official meeting minutes.
Answer using only the context below. Be concise.
If the context does not contain enough information, say so clearly.
When referencing decisions or action items, tie them to what appears in the minutes.

Context:
{context}

Question: {input}

Answer:"""
)


def answer_with_rag(db: Session, session_id: str, message: str, top_k: int) -> tuple[str, list[dict]]:
    session = db.get(ChatSession, session_id)
    if not session:
        raise ValueError("chat session not found")

    meeting = db.get(Meeting, session.meeting_id)
    if not meeting:
        raise ValueError("meeting not found")

    if meeting.index_status != IndexStatus.ready:
        err = (
            "Meeting minutes are not indexed yet. Upload a PDF or DOCX, then run indexing, "
            "and wait until status is READY."
        )
        db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
        db.add(
            ChatMessage(
                session_id=session_id,
                role=MessageRole.assistant,
                content=err,
                citations_json={"items": []},
            )
        )
        db.commit()
        return err, []

    chroma_path = latest_chroma_dir_for_meeting(session.meeting_id).resolve()
    llm = ChatOllama(
        model=settings.ollama_chat_model,
        base_url=settings.ollama_base_url,
        temperature=0,
    )

    if not chroma_path.exists():
        err = "No search index on disk for this meeting. Run indexing again after uploading minutes."
        db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
        db.add(
            ChatMessage(
                session_id=session_id,
                role=MessageRole.assistant,
                content=err,
                citations_json={"items": []},
            )
        )
        db.commit()
        return err, []

    embeddings = OllamaEmbeddings(
        model=settings.ollama_embed_model,
        base_url=settings.ollama_base_url,
    )

    citations: list[dict] = []
    answer: str

    try:
        vector_store = Chroma(
            persist_directory=str(chroma_path.resolve()),
            embedding_function=embeddings,
        )
        docs = vector_store.similarity_search(message, k=top_k)
        citations = [
            {
                "chunk_id": d.metadata.get("chunk_id"),
                "chunk_index": d.metadata.get("chunk_index"),
            }
            for d in docs
        ]
        context = "\n\n".join(d.page_content for d in docs)
        chain = MEETING_PROMPT | llm
        if not context.strip():
            result = chain.invoke({"context": "(No matching passages retrieved.)", "input": message})
        else:
            result = chain.invoke({"context": context, "input": message})
        answer = result.content if hasattr(result, "content") else str(result)
    except Exception as exc:
        try:
            chain = MEETING_PROMPT | llm
            result = chain.invoke(
                {
                    "context": f"(Retrieval failed: {exc!s}. Answering without retrieved passages.)",
                    "input": message,
                }
            )
            answer = result.content if hasattr(result, "content") else str(result)
        except Exception as exc2:
            answer = (
                "The meeting assistant could not reach Ollama. Ensure Ollama is running and models "
                f"`{settings.ollama_chat_model}` and `{settings.ollama_embed_model}` are available. "
                f"Details: {exc2!s}"
            )

    db.add(ChatMessage(session_id=session_id, role=MessageRole.user, content=message))
    db.add(
        ChatMessage(
            session_id=session_id,
            role=MessageRole.assistant,
            content=answer,
            citations_json={"items": citations},
        )
    )
    db.commit()
    return answer, citations
